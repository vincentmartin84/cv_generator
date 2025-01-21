import  os
from docx import  Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from functions import valided_file_name

#create new file
def create_file():
    cv_doc = Document()
    return cv_doc

#save the file
def save_file(document, name_file):

    try:
        save_doc= document.save(name_file )
        place= os.getcwd()
        print(f"Votre cv : {name_file} a bien été enregistré!")
        print(f"Emplacement du fichier: {place}")
    except Exception as E:
        print(f"Error:{e}")


#create new file
cv= create_file()


#create header Cv
def header_cv(document, contact, title):

    table = document.add_table(rows=1, cols=2)
    cell_contact = table.cell(0, 0)
    cell_title = table.cell(0, 1)

    # Write data in cell contact
    para_names= cell_contact.add_paragraph()
    run_names= para_names.add_run(contact["first_name"] + " " + contact["last_name"])

    para_tel= cell_contact.add_paragraph()
    run_tel = para_tel.add_run("tel: " + contact["phone"])

    para_email= cell_contact.add_paragraph()
    run_email= para_email.add_run("Email: " + contact["email"])

    para_residence=  cell_contact.add_paragraph()
    run_residence= para_residence.add_run("Domicile: " + contact["residence"])

    if contact["mention"]:
        para_mention= cell_contact.add_paragraph()
        run_mention = para_mention.add_run(contact["mention"])

    # Write data in cell title
    para_title= cell_title.add_paragraph()
    run_title= para_title.add_run(title["title"])

    # Styles title
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.small_caps = True
    run_title.alignment = WD_ALIGN_PARAGRAPH.CENTER


#trainings part
def training_part(document, training):

    #training part title
    training_title= document.add_heading("Formations", level=2)


    #write trainings
    for i, trn in enumerate(training,1):
        para_trainig= document.add_paragraph(style='List Bullet')
        run_label = para_trainig.add_run(trn["title"] + " ")
        run_school = para_trainig.add_run(trn["school"] + " ")
        run_date= para_trainig.add_run(trn["date_start"] + ":" + trn["date_end"])

        #styles tranings

#professionnal experiences
def pro_experiences_part(document, experience):
    #experiences title
    exp_title= document.add_heading("Expériences professionnelles : ", level=2)

    #write experiences
    for i, exp in enumerate(experience,1):
        para_exp= document.add_paragraph(style='List Bullet')
        run_title= para_exp.add_run(exp["title"] + " ")
        run_company= para_exp.add_run(exp["company"] + " ")
        run_date= para_exp.add_run(exp["date_start"] + " " + exp["date_end"] + " : ")
        run_summary= para_exp.add_run(exp["summary"])


#technical skills
def tech_skills(document, skills):

    #technical skills title
    skills_title = document.add_heading("Compétences techniques", level=2)

    #write skills
    para_skills= document.add_paragraph(skills)

#write professionnal skills
def pro_skills(document, skills):

    #professionnal skills title
    pro_title= document.add_heading("Compétences professionnelles", level=2)


    #write professionnal skills
    para_skills= document.add_paragraph(skills)



#personal project
def perso_projects(document, projects):

    #personal project title
    project_title = document.add_heading("Projets personnels", level=2)

    #write personal projects
    for i, prj in enumerate(projects,1):
        para_project = document.add_paragraph(style='List Bullet')
        run_title= para_project.add_run(prj["title"] + " : ")
        run_summary= para_project.add_run(prj["summary"] + "\n")
        run_link= para_project.add_run(prj["link"])


#file name
def file_name():

    print("\n Enregistrement du fichier \n")

    file_name= input("Saisir le nomdu fichier : ")

    if not valided_file_name(file_name):
        print("Le nom du fichier est incorrecte il ne doit contenir que des lettres, des chiffres et les caractères suivants: _ - ")
    return file_name